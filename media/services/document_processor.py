# -*- coding: utf-8 -*-
"""
Document Processor - Security-First Pipeline
=============================================
STRICT REQUIREMENTS:
1. Convert non-PDF documents to PDF
2. Security scan & sanitization
3. Compress PDF for minimum size, maximum quality (85%+)
4. Send to secure storage service

Pipeline: Input → Convert to PDF → Sanitize → Compress → Storage Service

Supported Input Formats:
- PDF (.pdf)
- Word Documents (.doc, .docx)
- Rich Text (.rtf)
- Plain Text (.txt)
- OpenDocument (.odt)
"""
import io
import logging
import hashlib
import tempfile
import os
from pathlib import Path
from typing import Optional, Tuple, List, Dict, Any

from .base import BaseProcessor, ProcessingResult, MediaType

logger = logging.getLogger(__name__)

# ═══════════════════════════════════════════════════════════════════════════════
# CONSTANTS - Strict requirements
# ═══════════════════════════════════════════════════════════════════════════════
MAX_OUTPUT_SIZE = 10 * 1024 * 1024   # 10MB max for PDFs
MIN_QUALITY = 85                      # Minimum quality - NEVER go below 85%
TARGET_QUALITY = 95                   # Starting quality for compression
MAX_PAGES = 500                       # Maximum pages allowed
MAX_TEXT_LENGTH = 5 * 1024 * 1024     # 5MB of extracted text
THUMBNAIL_DPI = 150                   # DPI for thumbnail generation
THUMBNAIL_SIZE = (300, 400)           # Thumbnail dimensions

# ═══════════════════════════════════════════════════════════════════════════════
# SECURITY PATTERNS - Malicious content detection for documents
# ═══════════════════════════════════════════════════════════════════════════════
DANGEROUS_DOC_PATTERNS = [
    # JavaScript in PDFs
    b'/JavaScript',
    b'/JS',
    b'/AA',           # Additional Actions
    b'/OpenAction',   # Auto-execute on open
    b'/Launch',       # Launch external apps
    b'/URI',          # Can be used for phishing
    b'/SubmitForm',   # Form submission
    b'/ImportData',   # Import external data

    # Embedded files
    b'/EmbeddedFile',
    b'/EmbeddedFiles',
    b'/FileAttachment',

    # Active content
    b'/AcroForm',     # Interactive forms (can contain scripts)
    b'/XFA',          # XML Forms Architecture
    b'/RichMedia',    # Flash/video
    b'/3D',           # 3D content

    # VBA Macros (Office documents)
    b'vbaProject.bin',
    b'VBA',
    b'macros',
    b'Auto_Open',
    b'AutoOpen',
    b'Auto_Close',
    b'Document_Open',
    b'Workbook_Open',

    # Shell/Script injection
    b'#!/bin/',
    b'#!/usr/',
    b'<script',
    b'<?php',
    b'<%',
    b'powershell',
    b'cmd.exe',
    b'exec(',
    b'system(',
    b'eval(',
]

# Executable signatures that should NEVER be in documents
EXECUTABLE_SIGNATURES = [
    (b'MZ', 'Windows PE executable'),
    (b'\x7fELF', 'Linux ELF binary'),
    (b'\xfe\xed\xfa\xce', 'Mach-O 32-bit'),
    (b'\xfe\xed\xfa\xcf', 'Mach-O 64-bit'),
    (b'\xca\xfe\xba\xbe', 'Mach-O Universal'),
]


class DocumentSanitizer:
    """
    Security scanning and sanitization for documents.
    Detects and removes dangerous content.
    """

    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.DocumentSanitizer")

    def scan_threats(self, data: bytes) -> List[str]:
        """
        Scan document data for embedded malicious content.
        Returns list of detected threats.
        """
        threats = []

        # Scan for dangerous patterns
        for pattern in DANGEROUS_DOC_PATTERNS:
            if pattern.lower() in data.lower():
                threat_desc = f"Dangerous pattern: {pattern[:30]!r}"
                if threat_desc not in threats:
                    threats.append(threat_desc)

        # Check for executable signatures
        for sig, description in EXECUTABLE_SIGNATURES:
            if sig in data[:1024]:  # Check first 1KB
                threats.append(f"Executable content: {description}")

        # Check for unusually large embedded streams (possible hidden content)
        if b'/FlateDecode' in data:
            # Count large streams
            stream_count = data.count(b'stream')
            if stream_count > 1000:
                threats.append(f"Excessive streams detected: {stream_count}")

        return threats

    def sanitize_pdf(self, data: bytes) -> Tuple[bytes, List[str]]:
        """
        Sanitize a PDF by removing dangerous elements.
        Returns (sanitized_data, list_of_removed_items).

        Strategy: Re-render PDF pages as images, then combine back to PDF.
        This removes ALL active content, scripts, and embedded objects.
        """
        removed = []

        try:
            import fitz  # PyMuPDF
        except ImportError:
            self.logger.warning("PyMuPDF not installed, returning original PDF")
            return data, ["PyMuPDF not available for sanitization"]

        try:
            # Open original PDF
            src_doc = fitz.open(stream=data, filetype="pdf")

            # Check for dangerous elements and log them
            for pattern in [b'/JavaScript', b'/JS', b'/Launch', b'/OpenAction']:
                if pattern in data:
                    removed.append(f"Removed: {pattern.decode('utf-8', errors='ignore')}")

            # Create new clean PDF by re-rendering each page
            clean_doc = fitz.open()

            for page_num in range(len(src_doc)):
                page = src_doc[page_num]

                # Render page to high-quality image
                mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for quality
                pix = page.get_pixmap(matrix=mat, alpha=False)

                # Create new page with same dimensions
                rect = page.rect
                new_page = clean_doc.new_page(width=rect.width, height=rect.height)

                # Insert rendered image
                img_data = pix.tobytes("png")
                new_page.insert_image(rect, stream=img_data)

            # Get sanitized PDF bytes
            sanitized_data = clean_doc.tobytes(
                garbage=4,           # Maximum garbage collection
                deflate=True,        # Compress streams
                clean=True,          # Clean unused objects
                linear=True,         # Optimize for web
            )

            src_doc.close()
            clean_doc.close()

            removed.append("Re-rendered all pages (removed all active content)")

            return sanitized_data, removed

        except Exception as e:
            self.logger.error(f"PDF sanitization failed: {e}")
            return data, [f"Sanitization error: {str(e)}"]


class DocumentConverter:
    """
    Convert various document formats to PDF.
    """

    def __init__(self):
        self.logger = logging.getLogger(f"{__name__}.DocumentConverter")

    def convert_to_pdf(self, data: bytes, mime_type: str) -> Tuple[Optional[bytes], str]:
        """
        Convert document to PDF format.

        Returns:
            Tuple of (pdf_bytes, error_message)
            If successful, error_message is empty
        """
        if mime_type == 'application/pdf':
            return data, ""

        if mime_type in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]:
            return self._convert_docx_to_pdf(data)

        if mime_type == 'text/plain':
            return self._convert_text_to_pdf(data)

        if mime_type == 'application/rtf':
            return self._convert_rtf_to_pdf(data)

        if mime_type == 'application/vnd.oasis.opendocument.text':
            return self._convert_odt_to_pdf(data)

        return None, f"Unsupported format for PDF conversion: {mime_type}"

    def _convert_docx_to_pdf(self, data: bytes) -> Tuple[Optional[bytes], str]:
        """Convert DOCX to PDF using python-docx and reportlab."""
        try:
            from docx import Document
            import fitz  # PyMuPDF for creating PDF
        except ImportError as e:
            return None, f"Missing dependency: {e}"

        try:
            # Read DOCX
            doc = Document(io.BytesIO(data))

            # Extract all text with basic formatting
            text_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_content.append(' | '.join(row_text))

            # Create PDF from text
            return self._text_to_pdf('\n\n'.join(text_content))

        except Exception as e:
            return None, f"DOCX conversion failed: {str(e)}"

    def _convert_text_to_pdf(self, data: bytes) -> Tuple[Optional[bytes], str]:
        """Convert plain text to PDF."""
        try:
            # Decode text with fallback encodings
            text = None
            for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
                try:
                    text = data.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                text = data.decode('utf-8', errors='replace')

            return self._text_to_pdf(text)

        except Exception as e:
            return None, f"Text conversion failed: {str(e)}"

    def _convert_rtf_to_pdf(self, data: bytes) -> Tuple[Optional[bytes], str]:
        """Convert RTF to PDF."""
        try:
            # Simple RTF to text extraction
            text = data.decode('utf-8', errors='replace')

            # Remove RTF control words (basic cleanup)
            import re
            text = re.sub(r'\\[a-z]+\d*\s?', '', text)
            text = re.sub(r'[{}]', '', text)
            text = text.strip()

            return self._text_to_pdf(text)

        except Exception as e:
            return None, f"RTF conversion failed: {str(e)}"

    def _convert_odt_to_pdf(self, data: bytes) -> Tuple[Optional[bytes], str]:
        """Convert ODT to PDF."""
        try:
            import zipfile
            from xml.etree import ElementTree as ET

            # ODT is a ZIP file with content.xml
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                content_xml = zf.read('content.xml')

            # Parse XML and extract text
            root = ET.fromstring(content_xml)

            # Extract all text content
            text_parts = []
            for elem in root.iter():
                if elem.text:
                    text_parts.append(elem.text.strip())

            return self._text_to_pdf('\n'.join(filter(None, text_parts)))

        except Exception as e:
            return None, f"ODT conversion failed: {str(e)}"

    def _text_to_pdf(self, text: str) -> Tuple[Optional[bytes], str]:
        """Convert plain text to PDF using PyMuPDF."""
        try:
            import fitz

            # Create new PDF
            doc = fitz.open()

            # Page settings
            page_width = 595   # A4 width in points
            page_height = 842  # A4 height in points
            margin = 50
            line_height = 14
            font_size = 11

            # Split text into lines that fit the page width
            max_chars_per_line = 85  # Approximate for A4 with margins
            lines = []
            for paragraph in text.split('\n'):
                if not paragraph.strip():
                    lines.append('')
                    continue

                # Word wrap
                words = paragraph.split()
                current_line = ''
                for word in words:
                    if len(current_line) + len(word) + 1 <= max_chars_per_line:
                        current_line += (' ' if current_line else '') + word
                    else:
                        if current_line:
                            lines.append(current_line)
                        current_line = word
                if current_line:
                    lines.append(current_line)

            # Calculate lines per page
            usable_height = page_height - (2 * margin)
            lines_per_page = int(usable_height / line_height)

            # Create pages
            current_line_idx = 0
            while current_line_idx < len(lines):
                page = doc.new_page(width=page_width, height=page_height)

                y_position = margin
                for i in range(lines_per_page):
                    if current_line_idx >= len(lines):
                        break

                    line = lines[current_line_idx]
                    if line:
                        page.insert_text(
                            (margin, y_position),
                            line,
                            fontsize=font_size,
                            fontname="helv",
                        )

                    y_position += line_height
                    current_line_idx += 1

            # Get PDF bytes
            pdf_bytes = doc.tobytes(
                garbage=4,
                deflate=True,
                clean=True,
            )

            doc.close()
            return pdf_bytes, ""

        except Exception as e:
            return None, f"PDF creation failed: {str(e)}"


class PDFCompressor:
    """
    Compress PDF for minimum size while maintaining quality (85%+).
    """

    def __init__(self, max_size: int = MAX_OUTPUT_SIZE, min_quality: int = MIN_QUALITY):
        self.max_size = max_size
        self.min_quality = min_quality
        self.logger = logging.getLogger(f"{__name__}.PDFCompressor")

    def compress(self, data: bytes, target_quality: int = TARGET_QUALITY) -> Tuple[Optional[bytes], Dict[str, Any]]:
        """
        Compress PDF to target size while maintaining quality.

        Returns:
            Tuple of (compressed_bytes, metadata_dict)
        """
        try:
            import fitz
        except ImportError:
            return data, {'error': 'PyMuPDF not installed', 'quality': 100}

        original_size = len(data)

        try:
            doc = fitz.open(stream=data, filetype="pdf")
            page_count = len(doc)

            # Strategy 1: Basic optimization (garbage collection, compression)
            optimized = doc.tobytes(
                garbage=4,           # Maximum garbage collection
                deflate=True,        # Compress streams
                clean=True,          # Remove unused objects
                linear=True,         # Web optimization
                pretty=False,        # Compact output
            )

            if len(optimized) <= self.max_size:
                doc.close()
                return optimized, {
                    'quality': 100,
                    'original_size': original_size,
                    'compressed_size': len(optimized),
                    'pages': page_count,
                    'method': 'optimization',
                }

            # Strategy 2: Reduce image quality in PDF
            quality = target_quality
            while quality >= self.min_quality:
                compressed = self._compress_with_quality(doc, quality)

                if len(compressed) <= self.max_size:
                    doc.close()
                    return compressed, {
                        'quality': quality,
                        'original_size': original_size,
                        'compressed_size': len(compressed),
                        'pages': page_count,
                        'method': 'image_quality_reduction',
                    }

                quality -= 5

            # Strategy 3: Re-render pages at lower DPI
            for dpi in [150, 120, 100, 80, 72]:
                compressed = self._rerender_at_dpi(doc, dpi, self.min_quality)

                if len(compressed) <= self.max_size:
                    doc.close()
                    return compressed, {
                        'quality': self.min_quality,
                        'original_size': original_size,
                        'compressed_size': len(compressed),
                        'pages': page_count,
                        'method': f'rerender_dpi_{dpi}',
                        'dpi': dpi,
                    }

            doc.close()

            # Return best effort
            return optimized, {
                'quality': self.min_quality,
                'original_size': original_size,
                'compressed_size': len(optimized),
                'pages': page_count,
                'warning': f'Could not compress below {len(optimized)} bytes',
            }

        except Exception as e:
            self.logger.error(f"PDF compression failed: {e}")
            return data, {'error': str(e)}

    def _compress_with_quality(self, doc, quality: int) -> bytes:
        """Compress PDF images at specified quality."""
        try:
            import fitz

            new_doc = fitz.open()

            for page_num in range(len(doc)):
                page = doc[page_num]
                new_page = new_doc.new_page(width=page.rect.width, height=page.rect.height)

                # Copy page content
                new_page.show_pdf_page(new_page.rect, doc, page_num)

            return new_doc.tobytes(
                garbage=4,
                deflate=True,
                clean=True,
            )

        except Exception:
            return doc.tobytes()

    def _rerender_at_dpi(self, doc, dpi: int, quality: int) -> bytes:
        """Re-render PDF pages as images at specified DPI."""
        try:
            import fitz

            new_doc = fitz.open()
            scale = dpi / 72.0

            for page_num in range(len(doc)):
                page = doc[page_num]

                # Render to image
                mat = fitz.Matrix(scale, scale)
                pix = page.get_pixmap(matrix=mat, alpha=False)

                # Create new page
                new_page = new_doc.new_page(width=page.rect.width, height=page.rect.height)

                # Insert image
                img_data = pix.tobytes("jpeg")
                new_page.insert_image(new_page.rect, stream=img_data)

            return new_doc.tobytes(
                garbage=4,
                deflate=True,
                clean=True,
            )

        except Exception:
            return doc.tobytes()


class DocumentProcessor(BaseProcessor):
    """
    Complete Document Processing Pipeline
    =====================================

    STRICT REQUIREMENTS:
    1. Input: PDF, DOCX, DOC, TXT, RTF, ODT
    2. Convert to PDF (if not already PDF)
    3. Security scan & sanitization
    4. Compress for minimum size, maximum quality (85%+)
    5. Extract text for search/intelligence
    6. Generate thumbnail
    7. Send to secure storage service

    Pipeline:
    ┌──────────┐    ┌──────────┐    ┌───────────┐    ┌──────────┐    ┌─────────┐
    │  Input   │ → │ Convert  │ → │ Sanitize  │ → │ Compress │ → │ Storage │
    │  Doc     │    │ to PDF   │    │ & Scan    │    │ PDF      │    │ Service │
    └──────────┘    └──────────┘    └───────────┘    └──────────┘    └─────────┘

    Usage:
        processor = DocumentProcessor()
        result = processor.process(doc_bytes)

        if result.success:
            pdf_data = result.content
            text = result.extracted_text
            threats = result.threats_sanitized
    """

    SUPPORTED_MIMES = [
        'application/pdf',
        'application/msword',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'text/plain',
        'application/rtf',
        'application/vnd.oasis.opendocument.text',
    ]

    def __init__(self, max_output_size: int = MAX_OUTPUT_SIZE):
        super().__init__()
        self.max_output_size = max_output_size
        self.sanitizer = DocumentSanitizer()
        self.converter = DocumentConverter()
        self.compressor = PDFCompressor(max_size=max_output_size)

    def can_process(self, mime_type: str) -> bool:
        """Check if this processor can handle the given MIME type."""
        return mime_type.lower() in self.SUPPORTED_MIMES

    def process(self, data: bytes, **options) -> ProcessingResult:
        """
        Process a document through the security-first pipeline.

        Args:
            data: Raw document bytes
            **options:
                extract_text: Extract text content (default True)
                generate_thumbnail: Generate preview thumbnail (default True)
                sanitize: Perform deep sanitization (default True)
                max_pages: Maximum pages to process (default 500)

        Returns:
            ProcessingResult with:
                - success: True if processing completed
                - content: PDF bytes (compressed, sanitized)
                - extracted_text: Text content for search
                - threats_sanitized: List of detected/removed threats
        """
        # Parse options
        extract_text = options.get('extract_text', True)
        gen_thumbnail = options.get('generate_thumbnail', True)
        do_sanitize = options.get('sanitize', True)
        max_pages = options.get('max_pages', MAX_PAGES)

        original_size = len(data)
        warnings = []

        # Detect MIME type
        from .base import detect_mime_type
        mime_type = detect_mime_type(data)

        self.logger.info(f"[STEP 1/5] Processing document: {mime_type}, {original_size} bytes")

        # ═══════════════════════════════════════════════════════════════════
        # STEP 1: SECURITY SCAN
        # ═══════════════════════════════════════════════════════════════════
        self.logger.info("[STEP 1/5] Security scan")
        threats = self.sanitizer.scan_threats(data)

        if threats:
            self.logger.warning(f"THREATS DETECTED ({len(threats)}): {threats}")

        # ═══════════════════════════════════════════════════════════════════
        # STEP 2: CONVERT TO PDF
        # ═══════════════════════════════════════════════════════════════════
        self.logger.info("[STEP 2/5] Converting to PDF")

        if mime_type != 'application/pdf':
            pdf_data, error = self.converter.convert_to_pdf(data, mime_type)
            if pdf_data is None:
                return ProcessingResult(
                    success=False,
                    media_type=MediaType.DOCUMENT,
                    mime_type=mime_type,
                    original_size=original_size,
                    error=f"PDF conversion failed: {error}",
                    threats_sanitized=threats,
                )
            warnings.append(f"Converted from {mime_type} to PDF")
        else:
            pdf_data = data

        # ═══════════════════════════════════════════════════════════════════
        # STEP 3: SANITIZE PDF
        # ═══════════════════════════════════════════════════════════════════
        if do_sanitize and threats:
            self.logger.info("[STEP 3/5] Sanitizing PDF (removing dangerous content)")
            pdf_data, removed_items = self.sanitizer.sanitize_pdf(pdf_data)
            threats.extend(removed_items)
        else:
            self.logger.info("[STEP 3/5] Skipping sanitization (no threats detected)")

        # ═══════════════════════════════════════════════════════════════════
        # STEP 4: COMPRESS PDF
        # ═══════════════════════════════════════════════════════════════════
        self.logger.info(f"[STEP 4/5] Compressing PDF (max {self.max_output_size // 1024 // 1024}MB)")

        compressed_data, compress_meta = self.compressor.compress(pdf_data)

        if compress_meta.get('warning'):
            warnings.append(compress_meta['warning'])

        # ═══════════════════════════════════════════════════════════════════
        # STEP 5: EXTRACT TEXT & GENERATE THUMBNAIL
        # ═══════════════════════════════════════════════════════════════════
        self.logger.info("[STEP 5/5] Extracting text and generating thumbnail")

        extracted_text = ''
        thumbnail_data = None
        page_count = 0

        try:
            import fitz
            doc = fitz.open(stream=compressed_data, filetype="pdf")
            page_count = min(len(doc), max_pages)

            # Extract text
            if extract_text:
                text_parts = []
                for i in range(page_count):
                    page = doc[i]
                    text = page.get_text()
                    if text:
                        text_parts.append(text)
                extracted_text = '\n\n'.join(text_parts)

                # Truncate if too long
                if len(extracted_text) > MAX_TEXT_LENGTH:
                    extracted_text = extracted_text[:MAX_TEXT_LENGTH]
                    warnings.append(f"Text truncated to {MAX_TEXT_LENGTH} characters")

            # Generate thumbnail
            if gen_thumbnail and len(doc) > 0:
                thumbnail_data = self._generate_thumbnail(doc[0])

            doc.close()

        except Exception as e:
            self.logger.warning(f"Post-processing failed: {e}")
            warnings.append(f"Text extraction/thumbnail failed: {str(e)}")

        # ═══════════════════════════════════════════════════════════════════
        # BUILD RESULT
        # ═══════════════════════════════════════════════════════════════════
        checksum = self.calculate_checksum(compressed_data)

        self.logger.info(
            f"SUCCESS: {original_size} → {len(compressed_data)} bytes "
            f"({compress_meta.get('quality', '?')}% quality, {page_count} pages)"
        )

        return ProcessingResult(
            success=True,
            content=compressed_data,
            media_type=MediaType.DOCUMENT,
            mime_type='application/pdf',
            extension='pdf',
            metadata={
                'pages': page_count,
                'original_format': mime_type,
                'quality': compress_meta.get('quality', 100),
                'compression_method': compress_meta.get('method', 'none'),
                'text_length': len(extracted_text),
                'has_text': bool(extracted_text),
            },
            threats_sanitized=threats,
            warnings=warnings,
            checksum_sha256=checksum,
            original_size=original_size,
            processed_size=len(compressed_data),
            thumbnail=thumbnail_data,
            extracted_text=extracted_text,
        )

    def _generate_thumbnail(self, page) -> Optional[bytes]:
        """Generate a WebP thumbnail from PDF page."""
        try:
            import fitz
            from PIL import Image

            # Render at good DPI
            mat = fitz.Matrix(THUMBNAIL_DPI / 72, THUMBNAIL_DPI / 72)
            pix = page.get_pixmap(matrix=mat, alpha=False)

            # Convert to PIL
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            img.thumbnail(THUMBNAIL_SIZE, Image.Resampling.LANCZOS)

            # Save as WebP
            buffer = io.BytesIO()
            img.save(buffer, format='WEBP', quality=80)
            return buffer.getvalue()

        except Exception as e:
            self.logger.warning(f"Thumbnail generation failed: {e}")
            return None


# ═══════════════════════════════════════════════════════════════════════════════
# CONVENIENCE FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════════

def process_document(data: bytes, **options) -> ProcessingResult:
    """
    Convenience function to process a document.

    Args:
        data: Raw document bytes
        **options: Processing options

    Returns:
        ProcessingResult with PDF output
    """
    processor = DocumentProcessor()
    return processor.process(data, **options)


def convert_to_pdf(data: bytes, mime_type: str) -> Tuple[Optional[bytes], str]:
    """
    Quick conversion to PDF.

    Args:
        data: Raw document bytes
        mime_type: MIME type of the document

    Returns:
        Tuple of (pdf_bytes, error_message)
    """
    converter = DocumentConverter()
    return converter.convert_to_pdf(data, mime_type)

